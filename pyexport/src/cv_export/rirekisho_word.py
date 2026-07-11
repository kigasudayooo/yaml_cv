"""公式テンプレート(templates/rirekisho_template.docx)に data.yaml の内容を差し込んで
履歴書 .docx を生成する。

テンプレートは B5 サイズの厚生労働省系標準フォーマット相当で、以下の構成を持つ:
- tables[0]: 基本情報（氏名・生年月日・住所・連絡先など、8行×可変列）
- tables[1]: 学歴・職歴 テーブル（1ページ目、年/月/内容の3列 × 14データ行）
- tables[2]: 学歴・職歴の続き（8行）+ 資格・免許（6行）
- tables[3]: 志望の動機など（自由記述）
- tables[4]: 本人希望記入欄（自由記述）

行・列のインデックスはテンプレートの生XML構造を解析して確認したもの。
テンプレート自体の体裁（罫線・フォント・列幅）は変更せず、値の差し込みのみ行う。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell

from cv_export.styling import load_rirekisho_data, set_japanese_font

TEMPLATE_PATH = Path(__file__).parent / "templates" / "rirekisho_template.docx"
TEMPLATE_FONT = "ＭＳ 明朝"
VALUE_SIZE = 10.5


def _set_cell_text(cell: _Cell, text: str) -> None:
    """セルの最初の run のテキストを差し替える（既存の書式を極力維持する）。

    テンプレートの空欄セルは書式付きの空 run を持つことが多いため、run自体は
    使い回してテキストだけ書き換える。run が無い場合のみ新規追加する。
    """
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        run = paragraph.add_run(text)
        set_japanese_font(run, name=TEMPLATE_FONT, size=VALUE_SIZE)


def _append_line(cell: _Cell, text: str) -> None:
    """セル内の既存テキスト(ラベル)の下に、新しい段落として値を追記する。"""
    if not text:
        return
    p = cell.add_paragraph()
    run = p.add_run(text)
    set_japanese_font(run, name=TEMPLATE_FONT, size=VALUE_SIZE)


def _combined_history_entries(data: dict[str, Any]) -> list[tuple[str, str, str]]:
    entries: list[tuple[str, str, str]] = [("", "", "学歴")]
    for item in data.get("education", []) or []:
        entries.append(
            (str(item.get("year", "")), str(item.get("month", "")), str(item.get("value", "")))
        )
    entries.append(("", "", "職歴"))
    for item in data.get("experience", []) or []:
        entries.append(
            (str(item.get("year", "")), str(item.get("month", "")), str(item.get("value", "")))
        )
    entries.append(("", "", "以上"))
    return entries


def _licence_entries(data: dict[str, Any]) -> list[tuple[str, str, str]]:
    return [
        (str(item.get("year", "")), str(item.get("month", "")), str(item.get("value", "")))
        for item in (data.get("licences", []) or [])
    ]


def _fill_entry_rows(
    rows_cells: list[tuple[_Cell, _Cell, _Cell]], entries: list[tuple[str, str, str]]
) -> None:
    for (year_cell, month_cell, value_cell), (year, month, value) in zip(
        rows_cells, entries, strict=False
    ):
        _set_cell_text(year_cell, year)
        _set_cell_text(month_cell, month)
        _set_cell_text(value_cell, value)


def build_rirekisho_document(data: dict[str, Any]) -> DocumentObject:
    """公式テンプレートに履歴書データを差し込んで python-docx Document を組み立てる。

    Args:
        data: load_rirekisho_data で読み込んだ data.yaml 相当の dict。

    Returns:
        テンプレートに値を差し込んだ python-docx Document。
    """
    document = Document(str(TEMPLATE_PATH))

    basic = document.tables[0]
    rows = basic.rows

    # 各行の row.cells は gridSpan に応じて同一セルが繰り返し出現する
    # （例: 行0は 履歴書(span2)+日付(span1) で合計3、行1以降はラベル(span1)+
    # 値(span2)+右欄(span1) で合計4）。そのため「右端の欄」のインデックスは
    # 行によって 2 だったり 3 だったりする点に注意。
    _set_cell_text(rows[0].cells[2], str(data.get("date", "")))
    _set_cell_text(rows[1].cells[1], str(data.get("name_kana", "")))
    _set_cell_text(rows[2].cells[1], str(data.get("name", "")))
    _set_cell_text(rows[3].cells[1], str(data.get("birth_day", "")))
    _append_line(rows[3].cells[3], str(data.get("gender", "")))

    _set_cell_text(rows[4].cells[1], str(data.get("address_kana", "")))
    # テンプレートには現住所ブロックに「電話」欄が1つのみ（携帯電話欄は無い）。
    # 固定電話(tel)があればそれを、無ければ携帯電話(cell_phone)を入れる。
    tel_primary = str(data.get("tel", "")) or str(data.get("cell_phone", ""))
    _append_line(rows[4].cells[3], tel_primary)

    zip_ = str(data.get("address_zip", ""))
    addr = str(data.get("address", ""))
    _set_cell_text(rows[5].cells[1], f"〒{zip_}　{addr}" if zip_ or addr else "〒")
    _append_line(rows[5].cells[3], str(data.get("email", "")))

    _set_cell_text(rows[6].cells[1], str(data.get("address_kana2", "")))
    _append_line(rows[6].cells[3], str(data.get("tel2", "")))

    zip2 = str(data.get("address_zip2", ""))
    addr2 = str(data.get("address2", ""))
    _set_cell_text(rows[7].cells[1], f"〒{zip2}　{addr2}" if zip2 or addr2 else "〒")
    # data.yaml には連絡先専用のメールアドレス項目が無いため、連絡先ブロックの
    # E-mail欄は空のままにする（同じ email を重複表示しない）。

    history_entries = _combined_history_entries(data)

    table1_rows = [(r.cells[0], r.cells[1], r.cells[2]) for r in document.tables[1].rows[1:]]
    table2 = document.tables[2]
    table2_history_rows = [(r.cells[0], r.cells[1], r.cells[2]) for r in table2.rows[1:9]]
    table2_licence_rows = [(r.cells[0], r.cells[1], r.cells[2]) for r in table2.rows[10:16]]

    _fill_entry_rows(table1_rows + table2_history_rows, history_entries)
    _fill_entry_rows(table2_licence_rows, _licence_entries(data))

    motivation = str(data.get("motivation", "")).strip()
    hobby = str(data.get("hobby", "")).strip()
    if hobby:
        motivation = (
            f"{motivation}\n\n【趣味・特技】\n{hobby}" if motivation else f"【趣味・特技】\n{hobby}"
        )
    _set_cell_text(document.tables[3].rows[1].cells[0], motivation)

    _set_cell_text(document.tables[4].rows[1].cells[0], str(data.get("request", "")))

    return document


def generate_rirekisho_word(input_file: Path, output_file: Path) -> None:
    """data.yaml から履歴書 .docx を生成してファイルに保存する。

    Args:
        input_file: data.yaml へのパス。
        output_file: 出力する .docx ファイルパス。
    """
    data = load_rirekisho_data(input_file)
    document = build_rirekisho_document(data)
    document.save(str(output_file))
