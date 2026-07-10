"""cv.md のパース結果から職務経歴書 .docx を生成する。"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from cv_export.md_parser import CVDocument, InlineRun, parse_cv_markdown, parse_inline
from cv_export.styling import set_japanese_font

SECTION_HEADING_SIZE = 13
COMPANY_HEADING_SIZE = 12
PROJECT_HEADING_SIZE = 11
BLOCK_HEADING_SIZE = 10.5
BODY_SIZE = 10.5


def _add_runs(paragraph: Paragraph, runs: list[InlineRun], size: float = BODY_SIZE) -> None:
    for r in runs:
        if r["text"] == "":
            continue
        run = paragraph.add_run(r["text"])
        run.bold = r["bold"]
        set_japanese_font(run, size=size)


def _add_bottom_border(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_pr.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "8", qn("w:space"): "2", qn("w:color"): "444444"},
    )
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _add_section_heading(document: DocumentObject, title: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(f"■ {title}")
    run.bold = True
    set_japanese_font(run, size=SECTION_HEADING_SIZE)
    _add_bottom_border(heading)


def _add_bullet(document: DocumentObject, text: str, size: float = BODY_SIZE) -> None:
    p = document.add_paragraph(style="List Bullet")
    _add_runs(p, parse_inline(text), size=size)


def build_shokumu_document(cv: CVDocument, name: str = "") -> DocumentObject:
    """CVDocument から職務経歴書の Document オブジェクトを組み立てる。

    Args:
        cv: parse_cv_markdown が返す構造化データ。
        name: 氏名。data.yaml 等から渡せる場合に使用し、無ければ空欄。

    Returns:
        構築済みの python-docx Document。
    """
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Cm(2.0))

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("職務経歴書")
    title_run.bold = True
    set_japanese_font(title_run, size=18)

    meta_p = document.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today = dt.date.today().strftime("%Y年%m月%d日")
    meta_text = f"{today} 現在"
    if name:
        meta_text += f"　　{name}"
    meta_run = meta_p.add_run(meta_text)
    set_japanese_font(meta_run, size=10)

    for sec in cv["sections"]:
        _add_section_heading(document, sec["title"])

        if sec["type"] == "text":
            for paragraph_text in sec["paragraphs"]:
                p = document.add_paragraph()
                _add_runs(p, parse_inline(paragraph_text))

        elif sec["type"] == "list":
            for item in sec["items"]:
                _add_bullet(document, item)

        elif sec["type"] == "skills":
            for group in sec["groups"]:
                gp = document.add_paragraph()
                gp.paragraph_format.space_before = Pt(6)
                grun = gp.add_run(group["heading"])
                grun.bold = True
                set_japanese_font(grun, size=BLOCK_HEADING_SIZE)
                for item in group["items"]:
                    _add_bullet(document, item)

        elif sec["type"] == "career":
            for company in sec["companies"]:
                cp = document.add_paragraph()
                cp.paragraph_format.space_before = Pt(10)
                crun = cp.add_run(company["name"])
                crun.bold = True
                set_japanese_font(crun, size=COMPANY_HEADING_SIZE)
                if company["period"]:
                    prun = cp.add_run(f"　（{company['period']}）")
                    set_japanese_font(prun, size=BODY_SIZE)

                for meta in company["meta"]:
                    mp = document.add_paragraph()
                    mrun = mp.add_run(f"{meta['key']}: {meta['value']}")
                    set_japanese_font(mrun, size=BODY_SIZE)

                for item in company["items"]:
                    _add_bullet(document, item)

                for project in company["projects"]:
                    pp = document.add_paragraph()
                    pp.paragraph_format.space_before = Pt(6)
                    prun = pp.add_run(project["name"])
                    prun.bold = True
                    set_japanese_font(prun, size=PROJECT_HEADING_SIZE)

                    for meta in project["meta"]:
                        mp = document.add_paragraph()
                        mrun = mp.add_run(f"{meta['key']}: {meta['value']}")
                        set_japanese_font(mrun, size=BODY_SIZE)

                    for block in project["blocks"]:
                        bp = document.add_paragraph()
                        brun = bp.add_run(block["heading"])
                        brun.bold = True
                        set_japanese_font(brun, size=BLOCK_HEADING_SIZE)
                        for item in block["items"]:
                            _add_bullet(document, item)

    return document


def generate_shokumu_word(input_file: Path, output_file: Path, name: str = "") -> None:
    """cv.md から職務経歴書 .docx を生成してファイルに保存する。

    Args:
        input_file: cv.md へのパス。
        output_file: 出力する .docx ファイルパス。
        name: 氏名（任意）。
    """
    cv = parse_cv_markdown(input_file)
    document = build_shokumu_document(cv, name=name)
    document.save(str(output_file))
